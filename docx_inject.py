"""
docx_inject.py — Insert device photos into Unit Test Word document
===================================================================
For each SN folder found under SN_FOLDERS_ROOT:
  1. Find the matching table in the docx (S/N at row[5][col=3])
  2. Collect all .jpg/.jpeg/.png images from the SN folder
  3. Insert them into table row[12] (the empty image placeholder area)

Structure expected:
  SN_FOLDERS_ROOT/
    WZP30149NFC/
      WZP30149NFC/          ← subfolder with same name (or images at top level)
        IMG_8402.JPG
        IMG_8403.JPG
        photo.heic

Usage:
  python docx_inject.py
"""

import sys, io, os, glob, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Cm, Pt, Emu
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import pillow_heif
import tempfile
import atexit

# Register HEIF/HEIC opener with Pillow
pillow_heif.register_heif_opener()

# Track temp files for cleanup
_temp_files = []


# ─────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────
# Find the ORIGINAL docx (not the output)
_all_docx = glob.glob(r'')
DOCX_PATH = [d for d in _all_docx if 'unit_test_with_images' not in d][0]
SN_FOLDERS_ROOT  = r''           # Parent folder containing WZP* subfolders
OUTPUT_DOCX      = r''

# Image cell location in the table (0-based indices)
IMG_ROW = 12     # row 12 = the empty image placeholder
IMG_COL = 0      # col 0  = leftmost (merged across all 4 cols)

# Max image size inside the cell
MAX_IMG_WIDTH_CM = 14.0


# ─────────────────────────────────────────────
# Step 1: Discover SN folders and their images
# ─────────────────────────────────────────────
VALID_IMG_EXTS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.heic', '.heif'}


def convert_heic_to_jpg(heic_path: str) -> str:
    """
    Convert a HEIC/HEIF file to a temporary JPG.
    Returns the path to the temp JPG file.
    """
    img = PILImage.open(heic_path)
    # Create temp file that persists until script exit
    fd, tmp_path = tempfile.mkstemp(suffix='.jpg', prefix='heic_conv_')
    os.close(fd)
    img.convert('RGB').save(tmp_path, 'JPEG', quality=90)
    _temp_files.append(tmp_path)
    return tmp_path


def cleanup_temp_files():
    """Remove all temporary converted files."""
    for f in _temp_files:
        try:
            os.remove(f)
        except OSError:
            pass


# Auto-cleanup on script exit
atexit.register(cleanup_temp_files)

def find_sn_folders(root: str) -> dict:
    """
    Scan root directory for WZP* folders.
    Returns dict: { sn_string: [list_of_image_paths] }
    Handles both flat and double-nested folder structures.
    """
    sn_map = {}
    if not os.path.isdir(root):
        return sn_map

    for entry in os.listdir(root):
        if not entry.startswith('WZP'):
            continue
        entry_path = os.path.join(root, entry)
        if not os.path.isdir(entry_path):
            continue

        sn = entry.strip()
        images = []

        # Collect images from this folder and all subfolders
        for dirpath, dirnames, filenames in os.walk(entry_path):
            for fname in sorted(filenames):
                ext = os.path.splitext(fname)[1].lower()
                # Skip temp conversion files from previous runs
                if '_converted' in fname:
                    continue
                if ext in VALID_IMG_EXTS:
                    images.append(os.path.join(dirpath, fname))

        if images:
            sn_map[sn] = images

    return sn_map


# ─────────────────────────────────────────────
# Step 2: Build SN → table index mapping from docx
# ─────────────────────────────────────────────
def build_sn_table_map(doc: Document) -> dict:
    """
    Scan all tables. Return dict: { sn_string: table_index }
    S/N is at row[5][col=3] in each device table.
    """
    sn_to_table = {}
    for t_idx, table in enumerate(doc.tables):
        try:
            sn_cell = table.rows[5].cells[3]
            sn = sn_cell.text.strip()
            if sn and sn.startswith('WZP'):
                sn_to_table[sn] = t_idx
        except (IndexError, AttributeError):
            continue
    return sn_to_table


# ─────────────────────────────────────────────
# Step 3: Insert images into table cells
# ─────────────────────────────────────────────
def insert_images_into_cell(cell, image_paths: list, max_width_cm: float = MAX_IMG_WIDTH_CM):
    """
    Insert one or more images into a table cell.
    Each image gets its own paragraph, centered.
    Images are scaled to max_width_cm while preserving aspect ratio.
    """
    # Clear existing content in the cell
    for para in cell.paragraphs:
        for run in para.runs:
            run._element.getparent().remove(run._element)
        # Remove empty paragraphs except the first
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    inserted_count = 0
    for img_idx, img_path in enumerate(image_paths):
        try:
            # Auto-convert HEIC/HEIF to JPG
            ext = os.path.splitext(img_path)[1].lower()
            if ext in ('.heic', '.heif'):
                original_name = os.path.basename(img_path)
                img_path = convert_heic_to_jpg(img_path)
                print(f"      Converted {original_name} -> JPG")

            # Get image dimensions for proper aspect ratio
            with PILImage.open(img_path) as pil_img:
                img_w, img_h = pil_img.size

            # Calculate width, maintaining aspect ratio
            width = Cm(max_width_cm)

            # Use first paragraph for first image, add new paragraph for subsequent
            if img_idx == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()

            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run()
            run.add_picture(img_path, width=width)
            inserted_count += 1

        except Exception as e:
            print(f"      ERROR inserting {os.path.basename(img_path)}: {e}")

    return inserted_count


# ─────────────────────────────────────────────
# Main Execution
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("=" * 60)
    print(" docx_inject.py — Unit Test Image Injector")
    print("=" * 60)
    print()

    # ── Step 1: Find SN folders ──
    print(f"Scanning SN folders in: {SN_FOLDERS_ROOT}")
    sn_folders = find_sn_folders(SN_FOLDERS_ROOT)
    print(f"  SN folders with images: {len(sn_folders)}")
    for sn, imgs in sorted(sn_folders.items()):
        print(f"    {sn}: {len(imgs)} image(s)")
        for img in imgs:
            print(f"      {os.path.basename(img)}")

    if not sn_folders:
        print("\n  ERROR: No SN folders found. Place WZP* folders in:")
        print(f"    {SN_FOLDERS_ROOT}")
        sys.exit(1)

    # ── Step 2: Open docx and map tables ──
    print(f"\nOpening docx: {os.path.basename(DOCX_PATH)}")
    doc = Document(DOCX_PATH)
    sn_table_map = build_sn_table_map(doc)
    print(f"  Tables with S/N: {len(sn_table_map)}")

    # ── Step 3: Inject images ──
    print(f"\n{'─' * 60}")
    print("Injecting images...")
    print(f"{'─' * 60}")

    injected   = 0   # tables with images added
    img_count  = 0   # total images inserted
    not_found  = 0   # SN folders with no matching table
    no_folder  = 0   # tables with no SN folder

    for sn, images in sorted(sn_folders.items()):
        t_idx = sn_table_map.get(sn)
        if t_idx is None:
            print(f"  [SKIP] {sn}: no matching table in docx")
            not_found += 1
            continue

        table = doc.tables[t_idx]
        try:
            img_cell = table.rows[IMG_ROW].cells[IMG_COL]
        except IndexError:
            print(f"  [ERR ] {sn}: table {t_idx} has no row {IMG_ROW}")
            continue

        count = insert_images_into_cell(img_cell, images)
        if count > 0:
            injected += 1
            img_count += count
            print(f"  [OK  ] {sn} (table {t_idx:>3}): {count} image(s) inserted")
        else:
            print(f"  [FAIL] {sn} (table {t_idx:>3}): no images could be inserted")

    # Count tables with no SN folder
    for sn in sn_table_map:
        if sn not in sn_folders:
            no_folder += 1

    # ── Step 4: Save ──
    print(f"\n{'=' * 60}")
    print(f"  Tables injected  : {injected}")
    print(f"  Images inserted  : {img_count}")
    print(f"  SN without table : {not_found}")
    print(f"  Tables no folder : {no_folder} (need more SN folders synced)")
    print(f"{'=' * 60}")

    doc.save(OUTPUT_DOCX)
    print(f"\nSaved -> {OUTPUT_DOCX}")
    print("Done!")
