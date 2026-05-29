"""
Unit Test Image Injector v1.1.0 — Portable GUI App
====================================================
Inserts device photos from SN folders into the
Unit Test Word document, matching by Serial Number.

Supports: .jpg .jpeg .png .bmp .tiff .heic .heif
Auto-converts HEIC/HEIF (iPhone photos) to JPG.

Changelog:
  v1.2.0 — Dynamic interactive table cell selection
  v1.1.1 — Bug fixes and optimizations
  v1.1.0 — Customizable folder keyword, Stop button, version display
  v1.0.0 — Initial release

Usage: python app.py
"""

__version__ = "1.2.0"
__app_name__ = "Unit Test Image Injector"

import sys
import os
import re
import glob
import tempfile
import atexit
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path

# ─────────────────────────────────────────────
# Core Engine
# ─────────────────────────────────────────────
VALID_IMG_EXTS = {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.heic', '.heif'}
_temp_files = []


def cleanup_temp_files():
    for f in _temp_files:
        try:
            os.remove(f)
        except OSError:
            pass


atexit.register(cleanup_temp_files)


def convert_heic_to_jpg(heic_path: str) -> str:
    from PIL import Image as PILImage
    img = PILImage.open(heic_path)
    fd, tmp_path = tempfile.mkstemp(suffix='.jpg', prefix='heic_conv_')
    os.close(fd)
    img.convert('RGB').save(tmp_path, 'JPEG', quality=90)
    _temp_files.append(tmp_path)
    return tmp_path


def find_sn_folders(root: str, keyword: str = "WZP") -> dict:
    """
    Scan root directory for folders starting with `keyword`.
    Returns dict: { folder_name: [list_of_image_paths] }
    """
    sn_map = {}
    if not os.path.isdir(root):
        return sn_map
    keyword_upper = keyword.strip().upper()
    if not keyword_upper:
        return sn_map  # Empty keyword would match everything — reject
    for entry in os.listdir(root):
        if not entry.upper().startswith(keyword_upper):
            continue
        entry_path = os.path.join(root, entry)
        if not os.path.isdir(entry_path):
            continue
        sn = entry.strip()
        images = []
        for dirpath, dirnames, filenames in os.walk(entry_path):
            for fname in sorted(filenames):
                ext = os.path.splitext(fname)[1].lower()
                if '_converted' in fname:
                    continue
                if ext in VALID_IMG_EXTS:
                    images.append(os.path.join(dirpath, fname))
        if images:
            sn_map[sn] = images
    return sn_map


def run_injection(docx_path: str, sn_root: str, output_path: str,
                  keyword: str = "WZP",
                  sn_row_col: tuple = (5, 3), img_row_col: tuple = (12, 0),
                  log_callback=None, progress_callback=None,
                  cancel_event=None):
    """
    Main injection logic. Returns (injected, img_count, no_folder, errors).
    cancel_event: threading.Event — set to stop processing early.
    """
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from PIL import Image as PILImage

    def log(msg):
        if log_callback:
            log_callback(msg)

    def is_cancelled():
        return cancel_event and cancel_event.is_set()

    # ── Step 1: Find SN folders ──
    log(f"Scanning folders matching '{keyword}*'...")
    sn_folders = find_sn_folders(sn_root, keyword)
    log(f"  Found {len(sn_folders)} folder(s) with images")

    if not sn_folders:
        log(f"ERROR: No '{keyword}*' folders found in the selected directory.")
        return 0, 0, 0, [f"No {keyword}* folders found"]

    # Show summary (avoid flooding log with 600+ lines)
    sn_sorted = sorted(sn_folders.keys())
    if len(sn_sorted) <= 10:
        for sn in sn_sorted:
            imgs = sn_folders[sn]
            heic_count = sum(1 for i in imgs if os.path.splitext(i)[1].lower() in ('.heic', '.heif'))
            log(f"    {sn}: {len(imgs)} file(s)" + (f" ({heic_count} HEIC)" if heic_count else ""))
    else:
        for sn in sn_sorted[:3]:
            imgs = sn_folders[sn]
            log(f"    {sn}: {len(imgs)} file(s)")
        log(f"    ... ({len(sn_sorted) - 6} more) ...")
        for sn in sn_sorted[-3:]:
            imgs = sn_folders[sn]
            log(f"    {sn}: {len(imgs)} file(s)")

    if is_cancelled():
        log("\n[CANCELLED] Stopped by user.")
        return 0, 0, 0, ["Cancelled"]

    # ── Step 2: Open docx ──
    log(f"\nOpening docx: {os.path.basename(docx_path)}")
    doc = Document(docx_path)
    log(f"  Tables found: {len(doc.tables)}")

    # Build SN → table index map
    keyword_upper = keyword.strip().upper()
    sn_to_table = {}
    for t_idx, table in enumerate(doc.tables):
        try:
            sn_cell = table.rows[sn_row_col[0]].cells[sn_row_col[1]]
            sn = sn_cell.text.strip()
            if sn and sn.upper().startswith(keyword_upper):
                sn_to_table[sn] = t_idx
        except (IndexError, AttributeError):
            continue

    log(f"  Tables with matching S/N: {len(sn_to_table)}")

    if is_cancelled():
        log("\n[CANCELLED] Stopped by user.")
        return 0, 0, 0, ["Cancelled"]

    # ── Step 3: Inject ──
    log(f"\n{'─' * 50}")
    log("Injecting images...")
    log(f"{'─' * 50}")

    injected = 0
    img_count = 0
    no_folder = 0
    errors = []
    total = len(sn_folders)
    current = 0

    for sn, images in sorted(sn_folders.items()):
        if is_cancelled():
            log(f"\n[CANCELLED] Stopped by user at {current}/{total}.")
            break

        current += 1
        if progress_callback:
            progress_callback(current, total)

        t_idx = sn_to_table.get(sn)
        if t_idx is None:
            log(f"  [SKIP] {sn}: no matching table in docx")
            continue

        table = doc.tables[t_idx]
        try:
            img_cell = table.rows[img_row_col[0]].cells[img_row_col[1]]
        except IndexError:
            log(f"  [ERR ] {sn}: table {t_idx} has no row {img_row_col[0]}")
            errors.append(f"{sn}: no row {img_row_col[0]}")
            continue

        # Clear cell
        for para in img_cell.paragraphs:
            for run in para.runs:
                run._element.getparent().remove(run._element)
        while len(img_cell.paragraphs) > 1:
            p = img_cell.paragraphs[-1]._element
            p.getparent().remove(p)

        cell_img_count = 0
        for img_idx, img_path in enumerate(images):
            if is_cancelled():
                break
            try:
                ext = os.path.splitext(img_path)[1].lower()
                if ext in ('.heic', '.heif'):
                    orig = os.path.basename(img_path)
                    img_path = convert_heic_to_jpg(img_path)
                    log(f"      Converted {orig} -> JPG")

                width = Cm(14.0)
                if img_idx == 0:
                    para = img_cell.paragraphs[0]
                else:
                    para = img_cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_path, width=width)
                cell_img_count += 1
            except Exception as e:
                err_msg = f"{sn}/{os.path.basename(img_path)}: {e}"
                log(f"      ERROR: {err_msg}")
                errors.append(err_msg)

        if cell_img_count > 0:
            injected += 1
            img_count += cell_img_count
            log(f"  [OK  ] {sn} (table {t_idx:>3}): {cell_img_count} image(s)")

    for sn in sn_to_table:
        if sn not in sn_folders:
            no_folder += 1

    # ── Step 4: Save ──
    cancelled = is_cancelled()
    if injected > 0:
        log(f"\n{'═' * 50}")
        log(f"  Tables injected  : {injected}")
        log(f"  Images inserted  : {img_count}")
        log(f"  Tables no folder : {no_folder}")
        log(f"  Errors           : {len(errors)}")
        if cancelled:
            log(f"  Status           : PARTIAL (cancelled)")
        log(f"{'═' * 50}")

        doc.save(output_path)
        log(f"\nSaved -> {output_path}")
    else:
        log("\nNo images were injected. Output file not saved.")

    if cancelled:
        log("Process was cancelled by user.")
    else:
        log("Done!")

    return injected, img_count, no_folder, errors


class TableSelectionDialog(tk.Toplevel):
    def __init__(self, parent, docx_path, current_sn_rc, current_img_rc):
        super().__init__(parent)
        self.title("Configure Table Layout")
        self.geometry("900x600")
        self.configure(bg="#1a1a2e")
        self.transient(parent)
        self.grab_set()

        self.result = None
        self.sn_rc = current_sn_rc
        self.img_rc = current_img_rc
        self.mode = tk.StringVar(value="sn")
        self.COLORS = parent.COLORS

        # Header
        header = tk.Frame(self, bg=self.COLORS["bg"])
        header.pack(fill="x", padx=10, pady=10)
        
        tk.Label(header, text="Select Cell Locations", bg=self.COLORS["bg"], fg=self.COLORS["text"], font=("Segoe UI", 14, "bold")).pack(side="left")
        
        instructions = tk.Label(header, text="1. Select 'S/N' mode and click the Serial Number cell.\n2. Select 'Image' mode and click the Empty Image cell.\n3. Click Approve.", bg=self.COLORS["bg"], fg=self.COLORS["text_dim"], font=("Segoe UI", 10), justify="left")
        instructions.pack(side="left", padx=20)

        # Controls
        controls = tk.Frame(self, bg=self.COLORS["card"])
        controls.pack(fill="x", padx=10, pady=5)
        
        tk.Radiobutton(controls, text="Select S/N Cell (Blue)", variable=self.mode, value="sn", bg=self.COLORS["card"], fg="#00a8ff", selectcolor=self.COLORS["bg"], font=("Segoe UI", 10, "bold")).pack(side="left", padx=10, pady=5)
        tk.Radiobutton(controls, text="Select Image Cell (Green)", variable=self.mode, value="img", bg=self.COLORS["card"], fg="#00b894", selectcolor=self.COLORS["bg"], font=("Segoe UI", 10, "bold")).pack(side="left", padx=10, pady=5)
        
        tk.Button(controls, text="Approve", bg=self.COLORS["btn_run"], fg=self.COLORS["btn_run_fg"], font=("Segoe UI", 10, "bold"), command=self._approve).pack(side="right", padx=10, pady=5)
        tk.Button(controls, text="Cancel", bg=self.COLORS["btn_stop"], fg=self.COLORS["btn_stop_fg"], font=("Segoe UI", 10, "bold"), command=self.destroy).pack(side="right")

        # Table Grid
        container = tk.Frame(self, bg=self.COLORS["bg"])
        container.pack(fill="both", expand=True, padx=10, pady=10)
        
        canvas = tk.Canvas(container, bg=self.COLORS["input_bg"], highlightthickness=0)
        vsb = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
        hsb = ttk.Scrollbar(container, orient="horizontal", command=canvas.xview)
        self.grid_frame = tk.Frame(canvas, bg=self.COLORS["input_bg"])
        
        self.grid_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.grid_frame, anchor="nw")
        canvas.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")

        self.buttons = {}
        self._load_table(docx_path)

    def _load_table(self, docx_path):
        from docx import Document
        try:
            doc = Document(docx_path)
            if not doc.tables:
                tk.Label(self.grid_frame, text="No tables found in the document.", fg=self.COLORS["accent"], bg=self.COLORS["input_bg"]).pack()
                return
            
            table = doc.tables[0]
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()[:30]
                    if not text: text = "(Empty)"
                    
                    btn = tk.Button(self.grid_frame, text=text, width=15, height=2,
                                    bg=self.COLORS["card"], fg=self.COLORS["text"],
                                    activebackground=self.COLORS["card_alt"],
                                    command=lambda r=r_idx, c=c_idx: self._on_cell_click(r, c))
                    btn.grid(row=r_idx, column=c_idx, padx=1, pady=1, sticky="nsew")
                    self.buttons[(r_idx, c_idx)] = btn
            
            self._update_colors()
        except Exception as e:
            tk.Label(self.grid_frame, text=f"Error loading table:\n{e}", fg=self.COLORS["accent"], bg=self.COLORS["input_bg"]).pack()

    def _on_cell_click(self, r, c):
        if self.mode.get() == "sn":
            self.sn_rc = (r, c)
        else:
            self.img_rc = (r, c)
        self._update_colors()

    def _update_colors(self):
        for (r, c), btn in self.buttons.items():
            if (r, c) == self.sn_rc:
                btn.configure(bg="#00a8ff", fg="white")
            elif (r, c) == self.img_rc:
                btn.configure(bg="#00b894", fg="white")
            else:
                btn.configure(bg=self.COLORS["card"], fg=self.COLORS["text"])

    def _approve(self):
        self.result = (self.sn_rc, self.img_rc)
        self.destroy()


# ─────────────────────────────────────────────
# GUI Application
# ─────────────────────────────────────────────
class App(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title(f"{__app_name__} v{__version__}")
        self.geometry("850x720")
        self.minsize(720, 580)
        self.configure(bg="#1a1a2e")

        # ── Colors ──
        self.COLORS = {
            "bg":         "#1a1a2e",
            "card":       "#16213e",
            "card_alt":   "#0f3460",
            "accent":     "#e94560",
            "accent_ok":  "#00b894",
            "text":       "#eaeaea",
            "text_dim":   "#8892b0",
            "input_bg":   "#0d1b2a",
            "input_fg":   "#e0e0e0",
            "border":     "#233554",
            "btn_run":    "#00b894",
            "btn_run_fg": "#ffffff",
            "btn_stop":   "#e94560",
            "btn_stop_fg":"#ffffff",
            "btn_pick":   "#0f3460",
            "btn_pick_fg":"#e0e0e0",
        }

        style = ttk.Style(self)
        style.theme_use('clam')

        style.configure("Card.TFrame",  background=self.COLORS["card"])
        style.configure("BG.TFrame",    background=self.COLORS["bg"])
        style.configure("Title.TLabel",
                         background=self.COLORS["bg"],
                         foreground=self.COLORS["text"],
                         font=("Segoe UI", 18, "bold"))
        style.configure("Subtitle.TLabel",
                         background=self.COLORS["bg"],
                         foreground=self.COLORS["text_dim"],
                         font=("Segoe UI", 10))
        style.configure("Version.TLabel",
                         background=self.COLORS["bg"],
                         foreground=self.COLORS["card_alt"],
                         font=("Segoe UI", 9))
        style.configure("Status.TLabel",
                         background=self.COLORS["bg"],
                         foreground=self.COLORS["accent_ok"],
                         font=("Segoe UI", 10))
        style.configure("green.Horizontal.TProgressbar",
                         troughcolor=self.COLORS["input_bg"],
                         background=self.COLORS["accent_ok"],
                         thickness=8)

        self._running = False
        self._cancel_event = threading.Event()
        self.sn_row_col = (5, 3)
        self.img_row_col = (12, 0)
        self._build_ui()

    def _build_ui(self):
        C = self.COLORS

        # ── Header ──
        header = ttk.Frame(self, style="BG.TFrame")
        header.pack(fill="x", padx=20, pady=(18, 5))

        title_row = tk.Frame(header, bg=C["bg"])
        title_row.pack(fill="x")
        ttk.Label(title_row, text=__app_name__, style="Title.TLabel").pack(side="left")
        ttk.Label(title_row, text=f"v{__version__}", style="Version.TLabel").pack(side="left", padx=(8, 0), anchor="s")

        ttk.Label(header, text="Insert device photos from SN folders into Unit Test Word document",
                  style="Subtitle.TLabel").pack(anchor="w", pady=(2, 0))

        # ── Input Cards ──
        inputs_frame = ttk.Frame(self, style="BG.TFrame")
        inputs_frame.pack(fill="x", padx=20, pady=10)

        # Card 1: DOCX file
        self.docx_var = tk.StringVar()
        self._make_file_card(inputs_frame, "Word Document (.docx)",
                             self.docx_var, self._pick_docx, icon="📄", with_config=True)

        # Card 2: SN folder root + keyword
        self.sn_root_var = tk.StringVar()
        self.keyword_var = tk.StringVar(value="WZP")
        self._make_sn_card(inputs_frame)

        # Card 3: Output path
        self.output_var = tk.StringVar()
        self._make_file_card(inputs_frame, "Output File (.docx)",
                             self.output_var, self._pick_output, icon="💾")

        # ── Button Row ──
        btn_frame = ttk.Frame(self, style="BG.TFrame")
        btn_frame.pack(fill="x", padx=20, pady=(5, 5))

        self.run_btn = tk.Button(
            btn_frame,
            text="▶  Run Injection",
            font=("Segoe UI", 12, "bold"),
            bg=C["btn_run"], fg=C["btn_run_fg"],
            activebackground="#00a885", activeforeground="#fff",
            relief="flat", cursor="hand2", padx=24, pady=8,
            command=self._on_run
        )
        self.run_btn.pack(side="left")

        self.stop_btn = tk.Button(
            btn_frame,
            text="■  Stop",
            font=("Segoe UI", 12, "bold"),
            bg=C["btn_stop"], fg=C["btn_stop_fg"],
            activebackground="#c0392b", activeforeground="#fff",
            relief="flat", cursor="hand2", padx=24, pady=8,
            command=self._on_stop,
            state="disabled"
        )
        self.stop_btn.pack(side="left", padx=(10, 0))

        self.status_label = ttk.Label(btn_frame, text="Ready", style="Status.TLabel")
        self.status_label.pack(side="left", padx=15)

        # ── Progress Bar ──
        self.progress = ttk.Progressbar(self, style="green.Horizontal.TProgressbar",
                                         orient="horizontal", mode="determinate")
        self.progress.pack(fill="x", padx=20, pady=(0, 5))

        # ── Log Area ──
        log_frame = tk.Frame(self, bg=C["card"], bd=0, highlightthickness=1,
                             highlightbackground=C["border"])
        log_frame.pack(fill="both", expand=True, padx=20, pady=(0, 18))

        self.log_text = tk.Text(
            log_frame,
            bg=C["input_bg"], fg=C["input_fg"],
            font=("Cascadia Code", 9),
            insertbackground=C["text"],
            selectbackground=C["accent"],
            relief="flat", wrap="word",
            padx=10, pady=8,
            state="disabled"
        )
        scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side="right", fill="y")
        self.log_text.pack(fill="both", expand=True)

        self.log_text.tag_configure("ok",       foreground="#00b894")
        self.log_text.tag_configure("err",      foreground="#e94560")
        self.log_text.tag_configure("info",     foreground="#8892b0")
        self.log_text.tag_configure("title",    foreground="#e0e0e0", font=("Cascadia Code", 9, "bold"))
        self.log_text.tag_configure("cancel",   foreground="#f39c12")

        self._auto_detect_paths()

    def _make_file_card(self, parent, label_text, var, command, icon="📄", with_config=False):
        C = self.COLORS
        card = tk.Frame(parent, bg=C["card"], bd=0, highlightthickness=1,
                        highlightbackground=C["border"])
        card.pack(fill="x", pady=3)

        inner = tk.Frame(card, bg=C["card"])
        inner.pack(fill="x", padx=12, pady=8)

        top = tk.Frame(inner, bg=C["card"])
        top.pack(fill="x")
        tk.Label(top, text=f"{icon}  {label_text}",
                 bg=C["card"], fg=C["text"],
                 font=("Segoe UI", 10, "bold")).pack(side="left")

        btn = tk.Button(top, text="Browse",
                        font=("Segoe UI", 9),
                        bg=C["btn_pick"], fg=C["btn_pick_fg"],
                        activebackground=C["card_alt"],
                        relief="flat", cursor="hand2",
                        padx=12, pady=2,
                        command=command)
        btn.pack(side="right")
        
        if with_config:
            self.config_btn = tk.Button(top, text="⚙ Configure Layout",
                                        font=("Segoe UI", 9),
                                        bg=C["border"], fg=C["text"],
                                        activebackground=C["card_alt"],
                                        relief="flat", cursor="hand2",
                                        padx=12, pady=2,
                                        command=self._configure_layout)
            self.config_btn.pack(side="right", padx=(0, 10))

        path_label = tk.Label(inner, textvariable=var,
                              bg=C["card"], fg=C["text_dim"],
                              font=("Consolas", 9), anchor="w")
        path_label.pack(fill="x", pady=(4, 0))
        
        if with_config:
            self.layout_label = tk.Label(inner, text=f"Layout: S/N at R{self.sn_row_col[0]+1}C{self.sn_row_col[1]+1} | Image at R{self.img_row_col[0]+1}C{self.img_row_col[1]+1}",
                                         bg=C["card"], fg=C["accent_ok"], font=("Segoe UI", 9, "italic"), anchor="w")
            self.layout_label.pack(fill="x", pady=(2, 0))

    def _make_sn_card(self, parent):
        """Build the SN Folders card with an editable keyword field."""
        C = self.COLORS
        card = tk.Frame(parent, bg=C["card"], bd=0, highlightthickness=1,
                        highlightbackground=C["border"])
        card.pack(fill="x", pady=3)

        inner = tk.Frame(card, bg=C["card"])
        inner.pack(fill="x", padx=12, pady=8)

        # Top row: label + keyword entry + browse button
        top = tk.Frame(inner, bg=C["card"])
        top.pack(fill="x")

        tk.Label(top, text="📁  SN Folders Directory",
                 bg=C["card"], fg=C["text"],
                 font=("Segoe UI", 10, "bold")).pack(side="left")

        btn = tk.Button(top, text="Browse",
                        font=("Segoe UI", 9),
                        bg=C["btn_pick"], fg=C["btn_pick_fg"],
                        activebackground=C["card_alt"],
                        relief="flat", cursor="hand2",
                        padx=12, pady=2,
                        command=self._pick_sn_root)
        btn.pack(side="right")

        # Keyword entry inline
        kw_frame = tk.Frame(top, bg=C["card"])
        kw_frame.pack(side="right", padx=(0, 10))

        tk.Label(kw_frame, text="Keyword:",
                 bg=C["card"], fg=C["text_dim"],
                 font=("Segoe UI", 9)).pack(side="left")

        kw_entry = tk.Entry(kw_frame,
                            textvariable=self.keyword_var,
                            bg=C["input_bg"], fg=C["accent_ok"],
                            font=("Consolas", 10, "bold"),
                            insertbackground=C["text"],
                            relief="flat", width=12,
                            highlightthickness=1,
                            highlightbackground=C["border"],
                            highlightcolor=C["accent_ok"])
        kw_entry.pack(side="left", padx=(5, 0))

        tk.Label(kw_frame, text="*",
                 bg=C["card"], fg=C["accent_ok"],
                 font=("Consolas", 10, "bold")).pack(side="left")

        # Path display
        path_label = tk.Label(inner, textvariable=self.sn_root_var,
                              bg=C["card"], fg=C["text_dim"],
                              font=("Consolas", 9), anchor="w")
        path_label.pack(fill="x", pady=(4, 0))

    def _auto_detect_paths(self):
        """Leave paths empty — user always selects via Browse."""
        pass

    def _pick_docx(self):
        path = filedialog.askopenfilename(
            title="Select Unit Test Word Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if path:
            self.docx_var.set(path)
            self._configure_layout()

    def _configure_layout(self):
        docx = self.docx_var.get().strip()
        if not docx or not os.path.isfile(docx):
            messagebox.showerror("Error", "Please select a valid Word document (.docx) first.")
            return
            
        dialog = TableSelectionDialog(self, docx, self.sn_row_col, self.img_row_col)
        self.wait_window(dialog)
        
        if dialog.result:
            self.sn_row_col, self.img_row_col = dialog.result
            self.layout_label.configure(text=f"Layout: S/N at R{self.sn_row_col[0]+1}C{self.sn_row_col[1]+1} | Image at R{self.img_row_col[0]+1}C{self.img_row_col[1]+1}")
            self._log(f"Table layout configured: S/N={self.sn_row_col}, Image={self.img_row_col}")

    def _pick_sn_root(self):
        path = filedialog.askdirectory(
            title="Select folder containing serial number subfolders"
        )
        if path:
            self.sn_root_var.set(path)

    def _pick_output(self):
        path = filedialog.asksaveasfilename(
            title="Save output as",
            filetypes=[("Word Documents", "*.docx")],
            defaultextension=".docx",
            initialfile="unit_test_with_images.docx"
        )
        if path:
            self.output_var.set(path)

    def _log(self, msg):
        def _append():
            self.log_text.configure(state="normal")
            tag = None
            if msg.startswith("  [OK"):
                tag = "ok"
            elif "ERROR" in msg or "[ERR" in msg or "CRASH" in msg:
                tag = "err"
            elif "CANCELLED" in msg or "cancelled" in msg:
                tag = "cancel"
            elif msg.startswith("  [") or msg.startswith("    "):
                tag = "info"
            elif "═" in msg or "─" in msg:
                tag = "title"

            if tag:
                self.log_text.insert("end", msg + "\n", tag)
            else:
                self.log_text.insert("end", msg + "\n")
            self.log_text.see("end")
            self.log_text.configure(state="disabled")
        self.after(0, _append)

    def _update_progress(self, current, total):
        def _set():
            self.progress["maximum"] = total
            self.progress["value"] = current
            pct = int(100 * current / total) if total > 0 else 0
            self.status_label.configure(text=f"Processing... {current}/{total} ({pct}%)")
        self.after(0, _set)

    def _set_running(self, running: bool):
        def _update():
            self._running = running
            if running:
                self.run_btn.configure(state="disabled", bg="#555")
                self.stop_btn.configure(state="normal")
            else:
                self.run_btn.configure(state="normal", bg=self.COLORS["btn_run"])
                self.stop_btn.configure(state="disabled")
        self.after(0, _update)

    def _on_stop(self):
        if self._running:
            self._cancel_event.set()
            self.status_label.configure(text="Stopping...", foreground="#f39c12")
            self.stop_btn.configure(state="disabled")

    def _on_run(self):
        if self._running:
            return

        docx = self.docx_var.get().strip()
        sn_root = self.sn_root_var.get().strip()
        output = self.output_var.get().strip()
        keyword = self.keyword_var.get().strip()

        if not docx or not os.path.isfile(docx):
            messagebox.showerror("Error", "Please select a valid Word document (.docx)")
            return
        if not sn_root or not os.path.isdir(sn_root):
            messagebox.showerror("Error", "Please select a valid SN folders directory")
            return
        if not output:
            messagebox.showerror("Error", "Please specify an output file path")
            return
        if not keyword:
            messagebox.showerror("Error", "Please enter a folder keyword (e.g. WZP)")
            return

        # Clear
        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.configure(state="disabled")
        self.progress["value"] = 0
        self._cancel_event.clear()
        self._set_running(True)
        self.status_label.configure(text="Starting...", foreground=self.COLORS["accent_ok"])

        def worker():
            try:
                import pillow_heif
                pillow_heif.register_heif_opener()

                injected, img_count, no_folder, errors = run_injection(
                    docx, sn_root, output,
                    keyword=keyword,
                    sn_row_col=self.sn_row_col, img_row_col=self.img_row_col,
                    log_callback=self._log,
                    progress_callback=self._update_progress,
                    cancel_event=self._cancel_event
                )

                def finish():
                    self._set_running(False)
                    cancelled = self._cancel_event.is_set()
                    if cancelled:
                        self.status_label.configure(
                            text=f"Cancelled: {injected} tables saved before stop",
                            foreground="#f39c12")
                    elif errors:
                        self.status_label.configure(
                            text=f"Done: {injected} tables, {img_count} images, {len(errors)} error(s)",
                            foreground=self.COLORS["accent"])
                    else:
                        self.status_label.configure(
                            text=f"Done: {injected} tables, {img_count} images inserted",
                            foreground=self.COLORS["accent_ok"])
                    self.progress["value"] = self.progress["maximum"]
                self.after(0, finish)

            except Exception as e:
                import traceback
                self._log(f"\nCRASH: {e}")
                self._log(traceback.format_exc())
                def err_finish():
                    self._set_running(False)
                    self.status_label.configure(text=f"Error: {e}", foreground=self.COLORS["accent"])
                self.after(0, err_finish)

        t = threading.Thread(target=worker, daemon=True)
        t.start()


# ─────────────────────────────────────────────
# Entry Point
# ─────────────────────────────────────────────
if __name__ == "__main__":
    app = App()
    app.mainloop()
